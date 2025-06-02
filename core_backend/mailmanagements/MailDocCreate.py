from docx import Document
from django.core.files.storage import default_storage
from docx.shared import Inches
import os
from django.conf import settings
import jdatetime
from mailmanagements import models as md
def generate_word_document(mail_instance):
    # ایجاد یک سند جدید با سربرگ
    doc = Document(default_storage.path(mail_instance.header_file.name) if mail_instance.header_file else None)

    # افزودن تاریخ و شماره نامه
    jalali_date = jdatetime.date.fromgregorian(date=mail_instance.created_at.date())
    date_str = f"تاریخ: {jalali_date.strftime('%Y/%m/%d')}"
    doc.add_paragraph(date_str)
    doc.add_paragraph(f"شماره نامه: {mail_instance.mail_number}")

    # افزودن موضوع و محتوا
    doc.add_heading(mail_instance.subject, level=1)
    doc.add_paragraph(mail_instance.content)

    # افزودن فضای امضا
    doc.add_paragraph("\n\nامضا:")
    if mail_instance.signature_image:
        doc.add_picture(default_storage.path(mail_instance.signature_image.name), width=Inches(2))
    if mail_instance.sender:
        doc.add_paragraph(f"{mail_instance.sender.get_full_name()}")

    # ذخیره فایل
    output_path = os.path.join(settings.MEDIA_ROOT, f"generated_letters/{mail_instance.mail_number}.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path.replace(settings.MEDIA_ROOT, "").lstrip("/")

# ویو برای دانلود فایل
from django.http import FileResponse

def download_letter(request, mail_id):
    mail = md.Mail.objects.get(id=mail_id)
    file_path = generate_word_document(mail)
    return FileResponse(
        open(os.path.join(settings.MEDIA_ROOT, file_path), 'rb'),
        as_attachment=True,
        filename=f"{mail.mail_number}.docx"
    )